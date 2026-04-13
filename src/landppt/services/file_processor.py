"""
File Processing Service for LandPPT
Handles document upload and content extraction as specified in requires.md
"""

import os
import re
import asyncio
import logging
from typing import List, Dict, Any, Optional, Tuple
from pathlib import Path
import tempfile

# Document processing libraries
try:
    import docx
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from PIL import Image
    import pytesseract
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False

from ..api.models import FileUploadResponse

logger = logging.getLogger(__name__)


class FileProcessor:
    """Processes uploaded files and extracts content for PPT generation"""
    
    def __init__(self, default_file_processing_mode: str = "markitdown"):
        self.supported_formats = {
            '.pdf': self._process_pdf,
            '.docx': self._process_docx,
            '.txt': self._process_txt,
            '.md': self._process_markdown,
            '.jpg': self._process_image,
            '.jpeg': self._process_image,
            '.png': self._process_image,
        }
        self.default_file_processing_mode = default_file_processing_mode
        
        # Keywords for scenario detection
        self.scenario_keywords = {
            'tourism': ['旅游', '景点', '行程', '旅行', '观光', '度假', '酒店', '机票', '导游'],
            'education': ['教育', '学习', '课程', '培训', '知识', '科普', '儿童', '学生', '教学'],
            'analysis': ['分析', '数据', '统计', '研究', '报告', '调查', '图表', '趋势', '结论'],
            'history': ['历史', '古代', '文化', '传统', '遗产', '文物', '朝代', '事件', '人物'],
            'technology': ['技术', '科技', '创新', '数字', '智能', '人工智能', '互联网', '软件', '硬件'],
            'business': ['商业', '企业', '市场', '营销', '销售', '管理', '战略', '财务', '投资'],
            'general': ['介绍', '概述', '总结', '说明', '展示', '汇报', '演示', '分享']
        }
    
    async def process_file(
        self,
        file_path: str,
        filename: str,
        *,
        file_processing_mode: Optional[str] = None,
    ) -> FileUploadResponse:
        """Process uploaded file and extract content"""
        try:
            file_ext = Path(filename).suffix.lower()
            file_size = os.path.getsize(file_path)

            if file_ext == '.pdf':
                content = await self._process_pdf(file_path, file_processing_mode=file_processing_mode)
            else:
                if file_ext not in self.supported_formats:
                    raise ValueError(f"Unsupported file format: {file_ext}")

                # Process file based on type
                processor = self.supported_formats[file_ext]
                content = await processor(file_path)
            
            # Extract topics and suggest scenarios
            topics = self._extract_topics(content)
            scenarios = self._suggest_scenarios(content)
            
            return FileUploadResponse(
                filename=filename,
                size=file_size,
                type=file_ext,
                processed_content=content,
                extracted_topics=topics,
                suggested_scenarios=scenarios,
                message=f"文件 {filename} 处理成功，提取了 {len(content)} 个字符的内容"
            )
            
        except Exception as e:
            logger.error(f"Error processing file {filename}: {e}")
            raise ValueError(f"文件处理失败: {str(e)}")
    
    async def _process_docx(self, file_path: str) -> str:
        """Process DOCX file"""
        if not DOCX_AVAILABLE:
            raise ValueError("DOCX processing not available. Please install python-docx.")

        def _process_docx_sync(file_path: str) -> str:
            """同步处理DOCX文件（在线程池中运行）"""
            doc = Document(file_path)
            content_parts = []

            # Extract paragraphs
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    content_parts.append(text)

            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        content_parts.append(" | ".join(row_text))

            return "\n\n".join(content_parts)

        try:
            # 在线程池中执行文件处理以避免阻塞主服务
            import asyncio
            loop = asyncio.get_running_loop()
            return await loop.run_in_executor(None, _process_docx_sync, file_path)

        except Exception as e:
            logger.error(f"Error processing DOCX file: {e}")
            raise ValueError(f"DOCX 文件处理失败: {str(e)}")
    
    async def _process_pdf(self, file_path: str, *, file_processing_mode: Optional[str] = None) -> str:
        """Process PDF file (manual upload path keeps original behavior)."""
        mode = (file_processing_mode or self.default_file_processing_mode or "markitdown").strip() or "markitdown"
        if mode == "magic_pdf":
            mineru_content = await self._extract_pdf_with_mineru(file_path)
            if mineru_content:
                return mineru_content
            logger.warning("PDF处理模式=magic_pdf，但MinerU不可用或失败；将回退到 PyPDF2")

        return await self._extract_pdf_with_pypdf2(file_path)

    async def process_downloaded_pdf_with_priority(self, file_path: str) -> Tuple[str, str]:
        """
        For URL-downloaded PDFs only: MinerU first, then MarkItDown, then PyPDF2 fallback.
        Returns (content, used_mode).
        """
        mineru_content = await self._extract_pdf_with_mineru(file_path)
        if mineru_content:
            return mineru_content, "magic_pdf"

        markitdown_content = await self._extract_pdf_with_markitdown(file_path)
        if markitdown_content:
            return markitdown_content, "markitdown"

        logger.warning("URL下载PDF: MinerU 与 MarkItDown 均不可用或解析失败，回退到 PyPDF2")
        return await self._extract_pdf_with_pypdf2(file_path), "pypdf2"

    @staticmethod
    def _clean_markdown_content(content: str) -> str:
        """Normalize markdown output to keep formatting consistent."""
        if not content:
            return ""

        content = re.sub(r"\n{3,}", "\n\n", content)
        lines = [line.rstrip() for line in content.split("\n")]
        content = "\n".join(lines)
        content = re.sub(r"\n(#{1,6}\\s)", r"\n\n\\1", content)
        content = re.sub(r"(#{1,6}.*)\n([^#\n])", r"\\1\n\n\\2", content)
        return content.strip()

    async def _get_current_mineru_config(self) -> Tuple[Optional[str], Optional[str]]:
        """Load MinerU config for the current request user in async context."""
        current_uid = None
        try:
            from ..auth.request_context import current_user_id

            current_uid = current_user_id.get()
        except Exception:
            current_uid = None

        if current_uid is None:
            return None, None

        try:
            from .db_config_service import get_db_config_service

            cfg = get_db_config_service()
            mineru_api_key = await cfg.get_config_value("mineru_api_key", user_id=current_uid)
            mineru_base_url = await cfg.get_config_value("mineru_base_url", user_id=current_uid)
            return mineru_api_key, mineru_base_url
        except Exception as e:
            logger.debug(f"Failed to load MinerU config from DB (user_id={current_uid}): {e}")
            return None, None

    async def _extract_pdf_with_mineru(self, file_path: str) -> Optional[str]:
        """Try MinerU first; return None if unavailable or failed."""
        try:
            from summeryanyfile.core.mineru_api_client import MineruAPIClient
        except Exception as e:
            logger.debug(f"MinerU客户端不可用，跳过MinerU: {e}")
            return None

        mineru_api_key, mineru_base_url = await self._get_current_mineru_config()
        client = MineruAPIClient(api_key=mineru_api_key, base_url=mineru_base_url)
        try:
            if not client.is_available:
                logger.info("MinerU未配置，回退到 MarkItDown")
                return None

            logger.info("使用 MinerU API 提取 PDF 内容")
            md_content, _extra = await client.extract_markdown(file_path=file_path)
            cleaned = self._clean_markdown_content(md_content or "")
            if cleaned:
                return cleaned
            logger.warning("MinerU 返回内容为空，回退到 MarkItDown")
            return None
        except Exception as e:
            logger.warning(f"MinerU 解析失败，回退到 MarkItDown: {e}")
            return None
        finally:
            try:
                await client.close()
            except Exception:
                pass

    async def _extract_pdf_with_markitdown(self, file_path: str) -> Optional[str]:
        """Fallback parser when MinerU is unavailable."""
        try:
            from summeryanyfile.core.markitdown_converter import MarkItDownConverter
        except Exception as e:
            logger.warning(f"MarkItDown转换器不可用: {e}")
            return None

        def _convert_sync(path: str) -> str:
            converter = MarkItDownConverter(
                enable_plugins=False,
                use_magic_pdf=False,
                enable_cache=True,
                processing_mode="markitdown",
            )
            content, _encoding = converter.convert_file(path)
            # Keep consistency with other markdown cleaning paths.
            if hasattr(converter, "clean_markdown_content"):
                return converter.clean_markdown_content(content or "")
            return self._clean_markdown_content(content or "")

        try:
            loop = asyncio.get_running_loop()
            content = await loop.run_in_executor(None, _convert_sync, file_path)
            cleaned = (content or "").strip()
            if cleaned:
                logger.info("MarkItDown PDF 解析成功")
                return cleaned
            logger.warning("MarkItDown 返回内容为空")
            return None
        except Exception as e:
            logger.warning(f"MarkItDown PDF 解析失败: {e}")
            return None

    async def _extract_pdf_with_pypdf2(self, file_path: str) -> str:
        """Final fallback parser to avoid complete failure."""
        if not PDF_AVAILABLE:
            raise ValueError("PDF processing not available. Please install markitdown[all] or PyPDF2.")

        def _process_pdf_sync(path: str) -> str:
            content_parts = []
            with open(path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    text = page.extract_text()
                    if text and text.strip():
                        content_parts.append(text.strip())
            return "\n\n".join(content_parts)

        try:
            loop = asyncio.get_running_loop()
            return await loop.run_in_executor(None, _process_pdf_sync, file_path)
        except Exception as e:
            logger.error(f"Error processing PDF file with PyPDF2: {e}")
            raise ValueError(f"PDF 文件处理失败: {str(e)}")
    
    async def _process_txt(self, file_path: str) -> str:
        """Process TXT file"""
        def _process_txt_sync(file_path: str) -> str:
            """同步处理TXT文件（在线程池中运行）"""
            try:
                with open(file_path, 'r', encoding='utf-8') as file:
                    content = file.read()

                # Try different encodings if UTF-8 fails
                if not content.strip():
                    encodings = ['gbk', 'gb2312', 'latin1']
                    for encoding in encodings:
                        try:
                            with open(file_path, 'r', encoding=encoding) as file:
                                content = file.read()
                            if content.strip():
                                break
                        except:
                            continue

                return content.strip()
            except Exception as e:
                raise e

        try:
            # 在线程池中执行文件处理以避免阻塞主服务
            import asyncio
            loop = asyncio.get_running_loop()
            return await loop.run_in_executor(None, _process_txt_sync, file_path)

        except Exception as e:
            logger.error(f"Error processing TXT file: {e}")
            raise ValueError(f"TXT 文件处理失败: {str(e)}")
    
    async def _process_markdown(self, file_path: str) -> str:
        """Process Markdown file"""
        def _process_markdown_sync(file_path: str) -> str:
            """同步处理Markdown文件（在线程池中运行）"""
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
            return content.strip()

        try:
            # 在线程池中执行文件处理以避免阻塞主服务
            import asyncio
            loop = asyncio.get_running_loop()
            return await loop.run_in_executor(None, _process_markdown_sync, file_path)

        except Exception as e:
            logger.error(f"Error processing Markdown file: {e}")
            raise ValueError(f"Markdown 文件处理失败: {str(e)}")
    
    async def _process_image(self, file_path: str) -> str:
        """Process image file using OCR"""
        if not OCR_AVAILABLE:
            return "图片文件已上传，但 OCR 功能不可用。请安装 pytesseract 和 PIL 以启用文字识别。"

        def _process_image_sync(file_path: str) -> str:
            """同步处理图像文件（在线程池中运行）"""
            image = Image.open(file_path)

            # Perform OCR
            text = pytesseract.image_to_string(image, lang='chi_sim+eng')

            if not text.strip():
                return "图片文件已处理，但未能识别出文字内容。"

            return text.strip()

        try:
            # 在线程池中执行图像处理以避免阻塞主服务
            import asyncio
            loop = asyncio.get_running_loop()
            return await loop.run_in_executor(None, _process_image_sync, file_path)

        except Exception as e:
            logger.error(f"Error processing image file: {e}")
            return f"图片处理失败: {str(e)}"
    
    def _extract_topics(self, content: str) -> List[str]:
        """Extract potential topics from content"""
        if not content:
            return []
        
        topics = []
        
        # Extract sentences that might be topics (short, descriptive)
        sentences = re.split(r'[。！？\n]', content)
        
        for sentence in sentences:
            sentence = sentence.strip()
            # Look for topic-like sentences (10-50 characters, no common words)
            if 10 <= len(sentence) <= 50:
                # Avoid sentences with too many common words
                common_words = ['的', '是', '在', '有', '和', '与', '或', '但', '而', '了', '着', '过']
                common_count = sum(1 for word in common_words if word in sentence)
                
                if common_count <= 2:  # Not too many common words
                    topics.append(sentence)
        
        # Also extract potential titles (lines that are short and at the beginning)
        lines = content.split('\n')
        for i, line in enumerate(lines[:10]):  # Check first 10 lines
            line = line.strip()
            if 5 <= len(line) <= 30 and not line.endswith('：'):
                topics.append(line)
        
        # Remove duplicates and limit to top 10
        topics = list(dict.fromkeys(topics))[:10]
        
        return topics
    
    def _suggest_scenarios(self, content: str) -> List[str]:
        """Suggest appropriate scenarios based on content"""
        if not content:
            return ['general']
        
        content_lower = content.lower()
        scenario_scores = {}
        
        # Score each scenario based on keyword matches
        for scenario, keywords in self.scenario_keywords.items():
            score = 0
            for keyword in keywords:
                score += content_lower.count(keyword)
            
            if score > 0:
                scenario_scores[scenario] = score
        
        # Sort by score and return top scenarios
        sorted_scenarios = sorted(scenario_scores.items(), key=lambda x: x[1], reverse=True)
        
        # Return top 3 scenarios, or 'general' if no matches
        if sorted_scenarios:
            return [scenario for scenario, score in sorted_scenarios[:3]]
        else:
            return ['general']
    
    def get_supported_formats(self) -> List[str]:
        """Get list of supported file formats"""
        return list(self.supported_formats.keys())
    
    def validate_file(self, filename: str, file_size: int, max_size_mb: int = 100) -> Tuple[bool, str]:
        """Validate uploaded file"""
        file_ext = Path(filename).suffix.lower()
        
        # Check file extension
        if file_ext not in self.supported_formats:
            return False, f"不支持的文件格式: {file_ext}。支持的格式: {', '.join(self.supported_formats.keys())}"
        
        # Check file size
        max_size_bytes = max_size_mb * 1024 * 1024
        if file_size > max_size_bytes:
            return False, f"文件大小超过限制 ({max_size_mb}MB)。当前文件大小: {file_size / 1024 / 1024:.1f}MB"
        
        # Check specific format requirements
        if file_ext == '.docx' and not DOCX_AVAILABLE:
            return False, "DOCX 处理功能不可用，请联系管理员安装 python-docx"
        
        if file_ext == '.pdf' and not PDF_AVAILABLE:
            return False, "PDF 处理功能不可用，请联系管理员安装 PyPDF2"
        
        if file_ext in ['.jpg', '.jpeg', '.png'] and not OCR_AVAILABLE:
            return True, "图片文件可以上传，但文字识别功能不可用"
        
        return True, "文件验证通过"
    
    async def create_ppt_from_content(self, content: str, suggested_topic: str = None) -> Dict[str, Any]:
        """Create PPT generation request from processed content"""
        # Extract or suggest a topic
        if not suggested_topic:
            topics = self._extract_topics(content)
            suggested_topic = topics[0] if topics else "文档内容展示"
        
        # Suggest scenarios
        scenarios = self._suggest_scenarios(content)
        primary_scenario = scenarios[0] if scenarios else 'general'
        
        # Create a structured outline from content
        sections = self._create_content_sections(content)
        
        return {
            'topic': suggested_topic,
            'scenario': primary_scenario,
            'requirements': f"基于上传文档内容生成PPT，包含以下要点：\n{content}",
            'uploaded_content': content,
            'suggested_sections': sections,
            'language': 'zh'
        }
    
    def _create_content_sections(self, content: str) -> List[Dict[str, str]]:
        """Create structured sections from content"""
        sections = []

        # Split content into logical sections
        paragraphs = [p.strip() for p in content.split('\n\n') if p.strip()]

        # Create title slide
        sections.append({
            'type': 'title',
            'title': '文档内容展示',
            'subtitle': '基于上传文档生成'
        })

        # Create content slides (max 10)
        for i, paragraph in enumerate(paragraphs[:9]):
            if len(paragraph) > 50:  # Only use substantial paragraphs
                # Try to extract a title from the first sentence
                sentences = paragraph.split('。')
                title = sentences[0][:30] + '...' if len(sentences[0]) > 30 else sentences[0]

                sections.append({
                    'type': 'content',
                    'title': title or f'内容 {i+1}',
                    'content': paragraph[:300] + '...' if len(paragraph) > 300 else paragraph
                })

        # Add thank you slide
        sections.append({
            'type': 'thankyou',
            'title': '谢谢观看',
            'subtitle': '基于文档内容生成'
        })

        return sections

    def merge_multiple_files_to_markdown(self, files_content: List[Dict[str, str]]) -> str:
        """
        将多个文件的内容合并为一个完整的Markdown文档

        Args:
            files_content: 文件内容列表，每项包含 filename 和 content

        Returns:
            合并后的Markdown格式内容
        """
        if not files_content:
            return ""

        # 如果只有一个文件，直接返回其内容
        if len(files_content) == 1:
            return files_content[0]["content"]

        # 构建合并后的Markdown文档
        merged_parts = []

        # 添加文档标题
        merged_parts.append("# 合并文档内容\n")
        merged_parts.append(f"*本文档由 {len(files_content)} 个源文件合并生成*\n")
        merged_parts.append("---\n")

        # 添加目录
        merged_parts.append("## 📋 文档目录\n")
        for i, file_info in enumerate(files_content, 1):
            filename = file_info["filename"]
            merged_parts.append(f"{i}. [{filename}](#{self._sanitize_anchor(filename)})\n")
        merged_parts.append("\n---\n")

        # 添加每个文件的内容
        for i, file_info in enumerate(files_content, 1):
            filename = file_info["filename"]
            content = file_info["content"]

            # 添加文件标题（作为一级标题）
            merged_parts.append(f"\n## {i}. {filename} {{#{self._sanitize_anchor(filename)}}}\n")

            # 添加分隔线
            merged_parts.append("---\n")

            # 添加文件内容
            # 如果内容已经包含Markdown格式，保持原样
            # 否则将其格式化为段落
            if content.strip():
                # 检查是否已经是Markdown格式
                if self._is_markdown_formatted(content):
                    merged_parts.append(f"{content}\n")
                else:
                    # 将纯文本内容转换为段落
                    paragraphs = [p.strip() for p in content.split('\n\n') if p.strip()]
                    for paragraph in paragraphs:
                        merged_parts.append(f"{paragraph}\n\n")
            else:
                merged_parts.append("*（此文件无内容或内容提取失败）*\n")

            # 添加文件结束标记
            merged_parts.append("\n")

        # 添加文档结尾
        merged_parts.append("\n---\n")
        merged_parts.append("*文档结束*\n")

        return "".join(merged_parts)

    def _sanitize_anchor(self, text: str) -> str:
        """将文本转换为合法的Markdown锚点"""
        # 移除特殊字符，只保留字母数字和中文
        sanitized = re.sub(r'[^\w\u4e00-\u9fff-]', '-', text)
        # 移除多余的连字符
        sanitized = re.sub(r'-+', '-', sanitized)
        # 移除首尾的连字符
        sanitized = sanitized.strip('-')
        return sanitized.lower()

    def _is_markdown_formatted(self, content: str) -> bool:
        """检查内容是否已经是Markdown格式"""
        # 简单检查是否包含常见的Markdown语法
        markdown_indicators = [
            r'^#{1,6}\s',  # 标题
            r'\*\*.*\*\*',  # 粗体
            r'\*.*\*',  # 斜体
            r'^\s*[-*+]\s',  # 列表
            r'^\s*\d+\.\s',  # 有序列表
            r'\[.*\]\(.*\)',  # 链接
            r'```',  # 代码块
        ]

        for indicator in markdown_indicators:
            if re.search(indicator, content, re.MULTILINE):
                return True

        return False
